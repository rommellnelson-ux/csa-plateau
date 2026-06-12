from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Procedure_utilisation_securisee_CSA_Plateau.docx"

NAVY = "0D2B45"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GOLD = "C9A646"
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "F7EFD9"
LIGHT_RED = "FCE8E6"
LIGHT_GREEN = "E8F5E9"
MUTED = "667084"
WHITE = "FFFFFF"
BLACK = "1A1A2E"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, color=BLACK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run(r2)
    else:
        set_run(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_run(p.add_run(text))
    return p


def add_step(doc, number, title, detail):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [620, 8740])
    left, right = table.rows[0].cells
    set_cell_shading(left, NAVY)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(str(number)), size=12, color=WHITE, bold=True)
    p = right.paragraphs[0]
    set_run(p.add_run(title + " — "), bold=True, color=NAVY)
    set_run(p.add_run(detail))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc, title, text, fill=LIGHT_GOLD):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_run(p.add_run(title + "\n"), bold=True, color=NAVY)
    set_run(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in (
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(hp.add_run("CSA PLATEAU | PROCÉDURE D’UTILISATION SÉCURISÉE"), size=9, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("GUIDE OPÉRATIONNEL"), size=10, color=GOLD, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("Procédure d’utilisation sécurisée"), size=25, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    set_run(p.add_run("Logiciel de gestion médicale — CSA Garde Républicaine du Plateau"), size=14, color=DARK_BLUE)

    meta = doc.add_table(rows=3, cols=2)
    set_table_geometry(meta, [2700, 6660])
    for row, values in zip(meta.rows, [
        ("Adresse officielle", "https://rommellnelson-ux.github.io/csa-plateau/"),
        ("Version", "1.0 — 12 juin 2026"),
        ("Classification", "Usage interne — données médicales confidentielles"),
    ]):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_run(row.cells[0].paragraphs[0].add_run(values[0]), bold=True, color=NAVY)
        set_run(row.cells[1].paragraphs[0].add_run(values[1]))

    add_callout(
        doc,
        "Règle essentielle",
        "Chaque agent utilise uniquement son compte personnel. Les identifiants, mots de passe et codes MFA ne doivent jamais être partagés.",
        LIGHT_RED,
    )

    add_heading(doc, "1. Accès et responsabilités", 1)
    add_body(doc, "Les écrans sont attribués selon la fonction enregistrée dans Supabase. Toute tentative d’accès hors rôle est refusée et peut être tracée.")
    roles = doc.add_table(rows=1, cols=3)
    set_table_geometry(roles, [2200, 3000, 4160])
    headers = ["Profil", "Accès principal", "Limites"]
    for idx, text in enumerate(headers):
        set_cell_shading(roles.rows[0].cells[idx], NAVY)
        set_run(roles.rows[0].cells[idx].paragraphs[0].add_run(text), color=WHITE, bold=True)
    set_repeat_table_header(roles.rows[0])
    role_rows = [
        ("Accueil", "Réception, constantes, consultation, file", "Selon les permissions du compte"),
        ("AS", "Accueil patient, constantes, file du jour", "Pas de consultation ni de caisse"),
        ("Soins", "Consultation orientée, soins, observation", "Pas de fonctions chef/comptabilité"),
        ("Soins + Pharmacie", "Soins et gestion pharmacie", "Compte nominatif obligatoire"),
        ("Laboratoire", "Saisie des actes, résultats, feuilles CMU", "Pas d’accès aux autres services"),
        ("Comptabilité", "Caisse, clôture, rapports", "Pas d’accès clinique complet"),
        ("Médecin-Chef", "Tous les modules", "MFA obligatoire à chaque nouvelle session"),
        ("CDV", "Aucun accès actuellement", "Fonctionnalité non définie"),
    ]
    for ridx, values in enumerate(role_rows):
        cells = roles.add_row().cells
        if ridx % 2:
            for cell in cells:
                set_cell_shading(cell, "F7F9FC")
        for idx, text in enumerate(values):
            set_run(cells[idx].paragraphs[0].add_run(text), size=10.5)

    add_heading(doc, "2. Connexion", 1)
    add_step(doc, 1, "Ouvrir le site officiel", "Utiliser exclusivement l’adresse indiquée en première page et vérifier le cadenas HTTPS.")
    add_step(doc, 2, "S’identifier", "Saisir son email professionnel et son mot de passe personnel.")
    add_step(doc, 3, "Contrôler son profil", "Vérifier le nom et la fonction affichés en haut de l’écran avant toute saisie.")
    add_step(doc, 4, "Se déconnecter", "Cliquer sur Déconnexion dès que le poste est quitté, même pour une courte absence.")

    add_heading(doc, "3. MFA obligatoire du Médecin-Chef", 1)
    add_body(doc, "Lors de la première connexion après activation, l’application affiche un QR code.")
    for text in [
        "Installer Google Authenticator, Microsoft Authenticator ou Authy sur un téléphone sécurisé.",
        "Scanner le QR code affiché. Si le scan échoue, saisir la clé manuelle.",
        "Entrer le code à six chiffres généré par l’application, puis valider.",
        "Aux connexions suivantes, saisir le nouveau code à six chiffres après le mot de passe.",
        "Ne jamais photographier, envoyer ou imprimer le QR code ou la clé secrète.",
    ]:
        add_bullet(doc, text)
    add_callout(doc, "Perte du téléphone", "Ne pas créer un nouveau compte. Contacter l’administrateur Supabase pour retirer l’ancien facteur MFA après vérification d’identité.", LIGHT_RED)

    add_heading(doc, "4. Parcours patient standard", 1)
    steps = [
        ("Accueil", "Rechercher d’abord le patient. Créer un dossier uniquement s’il n’existe pas."),
        ("Constantes", "Sélectionner le patient et mesurer les valeurs avant validation."),
        ("Consultation", "Renseigner motif, antécédents, ordonnance et orientations utiles."),
        ("Service orienté", "Le laboratoire, les soins ou la pharmacie recherchent le patient orienté."),
        ("Facturation", "Contrôler statut, droits et montant avant validation."),
        ("Clôture", "La comptabilité rapproche les transactions et la caisse physique."),
    ]
    for idx, (title, detail) in enumerate(steps, 1):
        add_step(doc, idx, title, detail)

    add_heading(doc, "5. Qualité des données cliniques", 1)
    add_body(doc, "L’application refuse les valeurs manifestement impossibles. En cas de doute, refaire la mesure plutôt que forcer une valeur.")
    values = doc.add_table(rows=1, cols=3)
    set_table_geometry(values, [2400, 2300, 4660])
    for idx, text in enumerate(["Mesure", "Plage acceptée", "Conduite"]):
        set_cell_shading(values.rows[0].cells[idx], NAVY)
        set_run(values.rows[0].cells[idx].paragraphs[0].add_run(text), color=WHITE, bold=True)
    set_repeat_table_header(values.rows[0])
    for measure, bounds, action in [
        ("Poids", "1 à 400 kg", "Vérifier l’unité et la balance"),
        ("Taille", "30 à 250 cm", "Saisir en centimètres"),
        ("Température", "30 à 45 °C", "Refaire la mesure si hors plage"),
        ("Pouls", "20 à 250 bpm", "Contrôler le capteur et le patient"),
        ("SpO2", "50 à 100 %", "Alerter immédiatement si valeur basse confirmée"),
        ("Tension", "Format 120/80", "Ne pas ajouter de texte dans le champ"),
    ]:
        cells = values.add_row().cells
        for idx, text in enumerate((measure, bounds, action)):
            set_run(cells[idx].paragraphs[0].add_run(text), size=10.5)
    add_callout(doc, "Correction d’une erreur", "Ne jamais recréer une fausse consultation pour masquer une erreur. Signaler l’enregistrement au Médecin-Chef afin que la correction soit tracée.", LIGHT_GOLD)

    add_heading(doc, "6. Confidentialité et sécurité", 1)
    for text in [
        "Ne consulter un dossier que pour les besoins directs de la prise en charge.",
        "Ne pas copier les données médicales dans WhatsApp, une messagerie personnelle ou un téléphone privé.",
        "Ne pas enregistrer les mots de passe dans un navigateur partagé.",
        "Ne pas exporter de CSV sans autorisation du Médecin-Chef.",
        "Verrouiller Windows ou se déconnecter avant de quitter le poste.",
        "Signaler immédiatement tout compte suspect, écran inhabituel ou perte d’appareil.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "7. Connexion, synchronisation et incidents", 1)
    add_body(doc, "Le voyant en haut de l’écran indique l’état de synchronisation.")
    status = doc.add_table(rows=1, cols=3)
    set_table_geometry(status, [1800, 2700, 4860])
    for idx, text in enumerate(["État", "Signification", "Action"]):
        set_cell_shading(status.rows[0].cells[idx], NAVY)
        set_run(status.rows[0].cells[idx].paragraphs[0].add_run(text), color=WHITE, bold=True)
    for row in [
        ("En ligne", "Données synchronisées", "Continuer normalement"),
        ("Sync…", "Envoi ou réception en cours", "Attendre avant de fermer"),
        ("Hors ligne", "Données conservées localement", "Ne pas vider le navigateur; reconnecter le réseau"),
        ("Erreur sync", "Échec d’échange avec Supabase", "Noter l’heure et prévenir le responsable"),
    ]:
        cells = status.add_row().cells
        for idx, text in enumerate(row):
            set_run(cells[idx].paragraphs[0].add_run(text), size=10.5)
    add_callout(doc, "Incident de confidentialité", "Déconnecter le poste sans supprimer les traces. Prévenir immédiatement le Médecin-Chef avec l’heure, le compte concerné et les faits observés.", LIGHT_RED)

    add_heading(doc, "8. Fin de poste — contrôle obligatoire", 1)
    checklist = [
        "Toutes les opérations en attente sont synchronisées.",
        "Les montants et stocks saisis ont été relus.",
        "Les impressions contenant des données patients sont récupérées.",
        "Aucun export médical ne reste sur le Bureau ou dans Téléchargements.",
        "La session est déconnectée.",
    ]
    for item in checklist:
        add_bullet(doc, "☐ " + item)

    add_heading(doc, "9. Administration et sauvegardes", 1)
    add_body(doc, "Cette section est réservée au Médecin-Chef ou à l’administrateur technique.")
    for text in [
        "Créer un compte Auth distinct par agent et un profil csa_profiles correspondant.",
        "Désactiver immédiatement les profils des agents mutés ou sortis.",
        "Vérifier régulièrement les journaux Auth et le journal de traçabilité.",
        "Conserver une sauvegarde externe chiffrée; le forfait Supabase Free n’inclut pas les sauvegardes automatiques.",
        "Tester périodiquement la restauration et l’accès GitHub Pages.",
        "Examiner la table csa_data_corrections après toute correction historique.",
    ]:
        add_bullet(doc, text)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
